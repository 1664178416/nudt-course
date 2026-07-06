# -*- coding: utf-8 -*-
"""
UNHCR 跨境保护关系网络可视化作业：独特数据版 Python 脚本

本脚本读取公开数据并导出题目要求的两个文件：
  姓名-学号/可视化作品.png
  姓名-学号/描述性文档.docx

数据主题：国家之间的“来源国 -> 庇护国”跨境保护关系网络。
节点：国家/地区；边：来源国到庇护国的登记人口关系；
权重：refugees + asylum_seekers + oip（难民、寻求庇护者、其他需国际保护者）。

运行前安装：
  pip install pandas numpy matplotlib networkx basemap pyproj python-docx pycountry countryinfo

提交前请把 STUDENT_NAME 和 STUDENT_ID 改成自己的真实姓名与学号。
"""

from __future__ import annotations

import math
import textwrap
import urllib.request
import warnings
from pathlib import Path
from typing import Dict, Iterable, Tuple
from functools import lru_cache

warnings.filterwarnings("ignore")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patheffects as pe
from matplotlib import patches
from matplotlib.collections import LineCollection
from matplotlib.colors import LinearSegmentedColormap, Normalize
from matplotlib.cm import ScalarMappable
from matplotlib.font_manager import FontProperties, fontManager
import networkx as nx
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from mpl_toolkits.basemap import Basemap
from pyproj import Geod

try:
    import pycountry
except Exception:  # pragma: no cover
    pycountry = None

try:
    from countryinfo import CountryInfo
except Exception:  # pragma: no cover
    CountryInfo = None

# ========== 0. 基本设置：提交前改这里 ==========
STUDENT_NAME = "姓名"
STUDENT_ID = "学号"
OUT_DIR = Path(f"{STUDENT_NAME}-{STUDENT_ID}")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PNG = OUT_DIR / "可视化作品.png"
OUT_DOCX = OUT_DIR / "描述性文档.docx"

DATA_DIR = Path("unhcr_network_data")
DATA_DIR.mkdir(exist_ok=True)
POP_FILE = DATA_DIR / "population.csv"
COUNTRY_FILE = DATA_DIR / "countries.csv"
POP_URL = "https://raw.githubusercontent.com/rfordatascience/tidytuesday/master/data/2023/2023-08-22/population.csv"
COUNTRY_URL = "https://raw.githubusercontent.com/google/dspl/master/samples/google/canonical/countries.csv"

YEAR = 2022
MIN_EDGE_FOR_GRAPH = 1000
MAP_EDGE_LIMIT = 320
COMMUNITY_EDGE_MIN = 10000
SEED = 2026

# ========== 1. 字体 ==========
def pick_font() -> tuple[FontProperties, FontProperties, str]:
    candidates = [
        "Microsoft YaHei", "SimHei", "PingFang SC", "Hiragino Sans GB",
        "Noto Sans CJK SC", "Noto Sans CJK JP", "Source Han Sans SC",
        "WenQuanYi Micro Hei", "Arial Unicode MS", "DejaVu Sans",
    ]
    names = {f.name for f in fontManager.ttflist}
    for name in candidates:
        if name in names:
            return FontProperties(family=name), FontProperties(family=name, weight="bold"), name
    return FontProperties(), FontProperties(weight="bold"), "DejaVu Sans"

FONT, FONT_BOLD, FONT_NAME = pick_font()
plt.rcParams["font.sans-serif"] = [FONT_NAME, "DejaVu Sans", "Arial"]
plt.rcParams["font.family"] = "sans-serif"
plt.rcParams["axes.unicode_minus"] = False

# ========== 2. 数据下载与清洗 ==========
def ensure_data() -> None:
    urls = {POP_FILE: POP_URL, COUNTRY_FILE: COUNTRY_URL}
    for path, url in urls.items():
        if not path.exists() or path.stat().st_size < 1000:
            print(f"Downloading {path.name} ...", flush=True)
            urllib.request.urlretrieve(url, path)


def iso2_to_iso3(alpha2: str) -> str | None:
    if pycountry is None or not isinstance(alpha2, str):
        return None
    try:
        obj = pycountry.countries.get(alpha_2=alpha2.upper())
        return obj.alpha_3 if obj is not None else None
    except Exception:
        return None


CN_NAMES = {
    "SYR": "叙利亚", "VEN": "委内瑞拉", "AFG": "阿富汗", "UKR": "乌克兰",
    "SSD": "南苏丹", "MMR": "缅甸", "COD": "刚果（金）", "SDN": "苏丹",
    "SOM": "索马里", "CAF": "中非", "ERI": "厄立特里亚", "IRQ": "伊拉克",
    "NGA": "尼日利亚", "BDI": "布隆迪", "NIC": "尼加拉瓜", "TUR": "土耳其",
    "IRN": "伊朗", "COL": "哥伦比亚", "DEU": "德国", "USA": "美国",
    "PAK": "巴基斯坦", "PER": "秘鲁", "UGA": "乌干达", "RUS": "俄罗斯",
    "POL": "波兰", "BGD": "孟加拉国", "ETH": "埃塞俄比亚", "LBN": "黎巴嫩",
    "JOR": "约旦", "FRA": "法国", "TCD": "乍得", "KEN": "肯尼亚",
    "CMR": "喀麦隆", "CZE": "捷克", "CHL": "智利", "BRA": "巴西",
    "ECU": "厄瓜多尔", "CRI": "哥斯达黎加", "GBR": "英国", "MEX": "墨西哥",
    "CAN": "加拿大", "ESP": "西班牙", "ITA": "意大利", "ZAF": "南非",
}

REGION_CN = {
    "Africa": "非洲", "Asia": "亚洲", "Europe": "欧洲", "Americas": "美洲",
    "Oceania": "大洋洲", "Unknown": "其他",
}

# 处理 countryinfo 无法识别或 UNHCR 名称特殊的常见代码
REGION_FALLBACK = {
    "SYR": "Asia", "VEN": "Americas", "AFG": "Asia", "UKR": "Europe", "SSD": "Africa",
    "MMR": "Asia", "COD": "Africa", "SDN": "Africa", "SOM": "Africa", "CAF": "Africa",
    "ERI": "Africa", "IRQ": "Asia", "NGA": "Africa", "BDI": "Africa", "NIC": "Americas",
    "TUR": "Asia", "IRN": "Asia", "COL": "Americas", "DEU": "Europe", "USA": "Americas",
    "PAK": "Asia", "PER": "Americas", "UGA": "Africa", "RUS": "Europe", "POL": "Europe",
    "BGD": "Asia", "ETH": "Africa", "LBN": "Asia", "JOR": "Asia", "FRA": "Europe",
    "TCD": "Africa", "KEN": "Africa", "CMR": "Africa", "CZE": "Europe", "CHL": "Americas",
    "BRA": "Americas", "ECU": "Americas", "CRI": "Americas", "GBR": "Europe", "MEX": "Americas",
    "CAN": "Americas", "ESP": "Europe", "ITA": "Europe", "ZAF": "Africa", "CHN": "Asia",
}


@lru_cache(maxsize=None)
def region_of_iso3(iso3: str) -> str:
    iso3 = str(iso3)
    if iso3 in REGION_FALLBACK:
        return REGION_FALLBACK[iso3]
    if pycountry is None or CountryInfo is None:
        return "Unknown"
    try:
        obj = pycountry.countries.get(alpha_3=iso3)
        if obj is None:
            return "Unknown"
        return CountryInfo(obj.name).region() or "Unknown"
    except Exception:
        return "Unknown"


def label_name(iso3: str, source_name: str | None = None) -> str:
    return CN_NAMES.get(str(iso3), str(source_name)[:12] if source_name else str(iso3))


def read_data() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, nx.DiGraph]:
    population = pd.read_csv(POP_FILE)
    countries = pd.read_csv(COUNTRY_FILE)
    countries["iso3"] = countries["country"].apply(iso2_to_iso3)
    countries = countries.dropna(subset=["iso3", "latitude", "longitude"]).copy()
    countries = countries.rename(columns={"latitude": "lat", "longitude": "lon", "name": "country_name"})
    countries = countries[["iso3", "lat", "lon", "country_name"]].drop_duplicates("iso3")

    pop = population.copy()
    for col in ["refugees", "asylum_seekers", "oip"]:
        pop[col] = pd.to_numeric(pop[col], errors="coerce").fillna(0)
    pop["value"] = pop["refugees"] + pop["asylum_seekers"] + pop["oip"]
    pop = pop[
        (pop["coo_iso"].notna()) & (pop["coa_iso"].notna()) &
        (pop["coo_iso"] != pop["coa_iso"]) &
        (~pop["coo_iso"].isin(["UNK", "XXA"])) &
        (~pop["coa_iso"].isin(["UNK", "XXA"])) &
        (pop["value"] > 0)
    ].copy()

    edges = pop[pop["year"] == YEAR].copy()
    edges = edges.merge(countries.add_prefix("s_"), left_on="coo_iso", right_on="s_iso3", how="inner")
    edges = edges.merge(countries.add_prefix("t_"), left_on="coa_iso", right_on="t_iso3", how="inner")
    edges["source_cn"] = [label_name(i, n) for i, n in zip(edges["coo_iso"], edges["coo_name"])]
    edges["target_cn"] = [label_name(i, n) for i, n in zip(edges["coa_iso"], edges["coa_name"])]
    edges["source_region"] = edges["coo_iso"].apply(region_of_iso3).map(REGION_CN).fillna("其他")
    edges["target_region"] = edges["coa_iso"].apply(region_of_iso3).map(REGION_CN).fillna("其他")
    edges["log_value"] = np.log10(edges["value"] + 1)

    graph_edges = edges[edges["value"] >= MIN_EDGE_FOR_GRAPH].copy()
    G = nx.DiGraph()
    for r in graph_edges.itertuples(index=False):
        G.add_node(r.coo_iso, name_cn=r.source_cn, lat=r.s_lat, lon=r.s_lon, region=r.source_region)
        G.add_node(r.coa_iso, name_cn=r.target_cn, lat=r.t_lat, lon=r.t_lon, region=r.target_region)
        G.add_edge(r.coo_iso, r.coa_iso, weight=float(r.value))

    return pop, countries, edges, G

# ========== 3. 作图辅助 ==========
def add_panel_frame(fig, xywh, face="#071923", edge="#22465b", alpha=0.88, lw=1.0, z=-1):
    x, y, w, h = xywh
    p = patches.FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.006,rounding_size=0.012",
        transform=fig.transFigure,
        facecolor=face, edgecolor=edge, linewidth=lw, alpha=alpha, zorder=z,
    )
    fig.patches.append(p)


def style_panel(ax, title: str | None = None):
    ax.set_facecolor("#071923")
    for sp in ax.spines.values():
        sp.set_color("#24475c")
        sp.set_linewidth(0.9)
    ax.tick_params(colors="#9db8c9", labelsize=8)
    ax.grid(color="#24475c", alpha=0.25, linewidth=0.7)
    if title:
        ax.set_title(title, loc="left", pad=8, color="#edf7ff", fontsize=11, fontproperties=FONT_BOLD)


def fmt_people(x: float) -> str:
    if x >= 1_000_000:
        return f"{x/1_000_000:.1f}M"
    if x >= 1_000:
        return f"{x/1_000:.0f}k"
    return f"{x:.0f}"


def great_circle_segments(df: pd.DataFrame, m: Basemap, geod: Geod, npts: int = 24):
    segments, weights = [], []
    width = m.xmax - m.xmin
    for r in df.itertuples(index=False):
        try:
            pts = geod.npts(float(r.s_lon), float(r.s_lat), float(r.t_lon), float(r.t_lat), npts)
            lons = np.asarray([r.s_lon] + [p[0] for p in pts] + [r.t_lon], dtype=float)
            lats = np.asarray([r.s_lat] + [p[1] for p in pts] + [r.t_lat], dtype=float)
            x, y = map(np.asarray, m(lons, lats))
            if not (np.isfinite(x).all() and np.isfinite(y).all()):
                continue
            cuts = np.where(np.abs(np.diff(x)) > width * 0.45)[0] + 1
            xs = np.split(x, cuts)
            ys = np.split(y, cuts)
            for sx, sy in zip(xs, ys):
                if len(sx) > 1:
                    segments.append(np.column_stack([sx, sy]))
                    weights.append(float(r.value))
        except Exception:
            continue
    return segments, np.asarray(weights)


def draw_metric_card(fig, x, y, w, h, title, value, note):
    card = patches.FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0.006,rounding_size=0.012",
        transform=fig.transFigure, facecolor="#0a2130", edgecolor="#2a5369",
        linewidth=0.9, alpha=0.93, zorder=3,
    )
    fig.patches.append(card)
    fig.text(x + 0.012, y + h - 0.024, title, color="#a9c6d7", fontsize=8.5, fontproperties=FONT, zorder=4)
    fig.text(x + 0.012, y + 0.030, value, color="#f8d47a", fontsize=18, fontproperties=FONT_BOLD, zorder=4)
    fig.text(x + 0.012, y + 0.010, note, color="#7fa2b5", fontsize=7.4, fontproperties=FONT, zorder=4)

# ========== 4. 主图 ==========
def make_visualization(pop: pd.DataFrame, countries: pd.DataFrame, edges: pd.DataFrame, G: nx.DiGraph) -> None:
    bg = "#04111b"
    panel_bg = "#071923"
    grid_c = "#24475c"
    text_c = "#edf7ff"
    muted = "#9db8c9"
    orange = "#f39b45"
    cyan = "#4fd6e8"
    gold = "#f8d47a"
    magenta = "#d564a7"
    blue = "#4f8bd7"

    cmap_flow = LinearSegmentedColormap.from_list(
        "flow_gold_magenta", ["#335c86", "#39c5d7", "#9f6ce3", "#f36f90", "#f8d47a"]
    )
    cmap_heat = LinearSegmentedColormap.from_list(
        "heat_refugee", ["#0f2b42", "#25577b", "#6c5ccf", "#d7609d", "#f5be59"]
    )
    community_colors = ["#48d7df", "#f2c75c", "#d769a6", "#73c16b", "#8d76e6", "#f28b54", "#6cb6ff", "#c6e66a"]

    fig = plt.figure(figsize=(20, 14), dpi=180, facecolor=bg)

    # 顶部标题与指标卡
    fig.text(0.030, 0.955, "全球跨境保护关系网络图谱", color=text_c, fontsize=28, fontproperties=FONT_BOLD)
    fig.text(0.030, 0.927, "节点=国家/地区  ·  边=来源国 → 庇护国  ·  权重=难民 + 寻求庇护者 + 其他需国际保护者  ·  年份=2022",
             color=muted, fontsize=10.5, fontproperties=FONT)
    fig.text(0.030, 0.907, "数据来源：UNHCR Refugee Data Finder，经 TidyTuesday population.csv 整理；地图节点使用国家质心，仅表示关系结构，不表示真实迁徙路线。",
             color="#6f8fa0", fontsize=8.7, fontproperties=FONT)

    total = edges["value"].sum()
    graph_edges = edges[edges["value"] >= MIN_EDGE_FOR_GRAPH]
    node_count = len(set(graph_edges["coo_iso"]).union(graph_edges["coa_iso"]))
    edge_count = len(graph_edges)
    density = nx.density(G)
    top_origin = edges.groupby(["coo_iso", "source_cn"])["value"].sum().sort_values(ascending=False).head(1)
    top_asylum = edges.groupby(["coa_iso", "target_cn"])["value"].sum().sort_values(ascending=False).head(1)
    components = list(nx.weakly_connected_components(G))
    largest_comp = max(len(c) for c in components) if components else 0

    card_y = 0.915
    draw_metric_card(fig, 0.545, card_y, 0.083, 0.065, "网络节点", f"{node_count}", f"边≥{MIN_EDGE_FOR_GRAPH:,}")
    draw_metric_card(fig, 0.636, card_y, 0.083, 0.065, "有效关系", f"{edge_count}", "来源-庇护边")
    draw_metric_card(fig, 0.727, card_y, 0.083, 0.065, "总权重", fmt_people(total), "跨境保护人口")
    draw_metric_card(fig, 0.818, card_y, 0.083, 0.065, "最大连通", f"{largest_comp}", "弱连通分量")
    draw_metric_card(fig, 0.909, card_y, 0.063, 0.065, "密度", f"{density:.3f}", "稀疏网络")

    # 面板位置
    positions = {
        "map": [0.030, 0.485, 0.610, 0.395],
        "role": [0.660, 0.485, 0.310, 0.395],
        "source": [0.030, 0.285, 0.295, 0.180],
        "asylum": [0.345, 0.285, 0.295, 0.180],
        "matrix": [0.660, 0.285, 0.310, 0.180],
        "trend": [0.030, 0.070, 0.300, 0.190],
        "community": [0.350, 0.070, 0.320, 0.190],
        "corridor": [0.690, 0.070, 0.280, 0.190],
    }
    for pos in positions.values():
        add_panel_frame(fig, pos, face=panel_bg, edge="#22465b", alpha=0.92)

    # 4.1 主地图
    ax_map = fig.add_axes(positions["map"])
    ax_map.set_facecolor(panel_bg)
    ax_map.set_title("01  世界保护关系流向图（Top weighted corridors）", loc="left", pad=10,
                     color=text_c, fontsize=12, fontproperties=FONT_BOLD)
    m = Basemap(projection="robin", lon_0=10, resolution="c", ax=ax_map)
    m.drawmapboundary(fill_color="#06131d", linewidth=0)
    m.fillcontinents(color="#142536", lake_color="#06131d", zorder=0)
    m.drawcoastlines(color="#345168", linewidth=0.35, zorder=1)
    m.drawcountries(color="#263e52", linewidth=0.25, zorder=1)

    geod = Geod(ellps="WGS84")
    map_edges = edges.nlargest(MAP_EDGE_LIMIT, "value")
    segs, w = great_circle_segments(map_edges, m, geod, npts=18)
    if len(segs):
        norm = Normalize(vmin=np.log10(max(1, w.min())), vmax=np.log10(w.max()))
        lc = LineCollection(segs, cmap=cmap_flow, norm=norm, alpha=0.35, zorder=2)
        lc.set_array(np.log10(w))
        lc.set_linewidths(0.25 + 2.6 * (np.log10(w) - np.log10(w.min())) / (np.log10(w.max()) - np.log10(w.min()) + 1e-9))
        ax_map.add_collection(lc)
        sm = ScalarMappable(cmap=cmap_flow, norm=norm)
        cax = fig.add_axes([positions["map"][0] + 0.018, positions["map"][1] + 0.030, 0.105, 0.012])
        cb = plt.colorbar(sm, cax=cax, orientation="horizontal")
        cb.ax.tick_params(labelsize=6.5, colors=muted, length=0)
        cb.set_ticks([norm.vmin, (norm.vmin + norm.vmax) / 2, norm.vmax])
        cb.set_ticklabels(["低", "中", "高"])
        cb.outline.set_edgecolor("#31546a")
        fig.text(positions["map"][0] + 0.018, positions["map"][1] + 0.047, "关系权重（log）", color=muted, fontsize=7.5, fontproperties=FONT)

    in_strength = edges.groupby("coa_iso")["value"].sum()
    out_strength = edges.groupby("coo_iso")["value"].sum()
    nodes = pd.DataFrame({"iso3": sorted(set(in_strength.index).union(out_strength.index))})
    nodes["in_strength"] = nodes["iso3"].map(in_strength).fillna(0)
    nodes["out_strength"] = nodes["iso3"].map(out_strength).fillna(0)
    nodes = nodes.merge(countries, on="iso3", how="inner")
    x, y = m(nodes["lon"].values, nodes["lat"].values)
    max_in = np.log10(nodes["in_strength"].max() + 1)
    max_out = np.log10(nodes["out_strength"].max() + 1)
    s_in = 18 + 115 * np.log10(nodes["in_strength"] + 1) / max_in
    s_out = 10 + 82 * np.log10(nodes["out_strength"] + 1) / max_out
    ax_map.scatter(x, y, s=s_in, c=cyan, alpha=0.33, edgecolor="none", zorder=3)
    ax_map.scatter(x, y, s=s_out, c=orange, alpha=0.60, edgecolor="#ffd08b", linewidth=0.25, zorder=4)

    label_codes = list(pd.Index(out_strength.sort_values(ascending=False).head(8).index).union(
        pd.Index(in_strength.sort_values(ascending=False).head(8).index)
    ))
    label_nodes = nodes[nodes["iso3"].isin(label_codes)].copy()
    offsets = {
        "SYR": (9, -8), "TUR": (9, 8), "AFG": (7, -10), "IRN": (10, 7), "VEN": (-34, -3),
        "COL": (-20, 8), "UKR": (5, 8), "DEU": (6, -8), "USA": (-20, 7), "PAK": (8, -12),
        "PER": (-18, -9), "UGA": (8, -8), "RUS": (8, 8), "POL": (-18, 8), "BGD": (8, 8),
        "LBN": (8, 8), "JOR": (8, -8), "SDN": (8, -8),
    }
    for r in label_nodes.itertuples(index=False):
        xx, yy = m(float(r.lon), float(r.lat))
        dx, dy = offsets.get(r.iso3, (6, 6))
        ax_map.text(xx + dx * 80000, yy + dy * 80000, f"{CN_NAMES.get(r.iso3, r.iso3)}\n{r.iso3}",
                    color="#f6fbff", fontsize=7.4, fontproperties=FONT_BOLD, zorder=5,
                    path_effects=[pe.withStroke(linewidth=2.0, foreground="#06131d")])
    ax_map.text(0.015, 0.055, "○ 蓝色面积：接收强度    ● 橙色面积：来源强度    线宽/颜色：关系权重",
                transform=ax_map.transAxes, color=muted, fontsize=8, fontproperties=FONT)
    ax_map.set_axis_off()

    # 4.2 角色矩阵
    ax_role = fig.add_axes(positions["role"])
    style_panel(ax_role, "02  国家角色矩阵：来源压力 × 接收承载")
    role_nodes = nodes.copy()
    role_nodes["total"] = role_nodes["in_strength"] + role_nodes["out_strength"]
    role_nodes["role"] = np.select(
        [
            (role_nodes["out_strength"] > role_nodes["in_strength"] * 3) & (role_nodes["out_strength"] > 100000),
            (role_nodes["in_strength"] > role_nodes["out_strength"] * 3) & (role_nodes["in_strength"] > 100000),
            (role_nodes["in_strength"] > 100000) & (role_nodes["out_strength"] > 100000),
        ],
        ["来源型", "庇护型", "双重节点"],
        default="一般节点",
    )
    role_color = {"来源型": orange, "庇护型": cyan, "双重节点": magenta, "一般节点": "#7d9aaa"}
    for role, sub in role_nodes.groupby("role"):
        ax_role.scatter(np.log10(sub["out_strength"] + 1), np.log10(sub["in_strength"] + 1),
                        s=18 + 120 * np.log10(sub["total"] + 1) / np.log10(role_nodes["total"].max() + 1),
                        c=role_color[role], alpha=0.68, label=role, edgecolors="white", linewidths=0.25)
    lim = [0, max(np.log10(role_nodes["out_strength"].max()+1), np.log10(role_nodes["in_strength"].max()+1)) + 0.2]
    ax_role.plot(lim, lim, color="#4a6b7c", lw=1.0, ls="--", alpha=0.5)
    ax_role.set_xlim(lim); ax_role.set_ylim(lim)
    ax_role.set_xlabel("来源强度 log10", color=muted, fontsize=8, fontproperties=FONT)
    ax_role.set_ylabel("庇护强度 log10", color=muted, fontsize=8, fontproperties=FONT)
    ax_role.legend(loc="lower right", frameon=True, facecolor="#0b2433", edgecolor="#2a5369",
                   fontsize=8, labelcolor=muted)
    # 标签
    key_labels = list(set(list(out_strength.sort_values(ascending=False).head(8).index) + list(in_strength.sort_values(ascending=False).head(8).index)))
    for r in role_nodes[role_nodes["iso3"].isin(key_labels)].itertuples(index=False):
        ax_role.text(np.log10(r.out_strength + 1) + 0.03, np.log10(r.in_strength + 1) + 0.03,
                     CN_NAMES.get(r.iso3, r.iso3), color=text_c, fontsize=7.3, fontproperties=FONT_BOLD,
                     path_effects=[pe.withStroke(linewidth=1.8, foreground=panel_bg)])
    ax_role.text(0.03, 0.96, "右上：既有来源又承担庇护；右下：主要来源国；左上：主要庇护国",
                 transform=ax_role.transAxes, color="#7fa2b5", fontsize=7.5, va="top", fontproperties=FONT)

    # 4.3 Top 来源国
    ax_src = fig.add_axes(positions["source"])
    style_panel(ax_src, "03  主要来源国（出边权重）")
    src_top = edges.groupby(["coo_iso", "source_cn"])["value"].sum().sort_values(ascending=False).head(10).reset_index()
    y_pos = np.arange(len(src_top))[::-1]
    colors_src = [orange if i < 4 else "#b76557" for i in range(len(src_top))]
    ax_src.barh(y_pos, src_top["value"][::-1] / 1e6, color=colors_src[::-1], alpha=0.86)
    ax_src.set_yticks(y_pos)
    ax_src.set_yticklabels([f"{a} {b}" for a, b in zip(src_top["coo_iso"], src_top["source_cn"])][::-1],
                           fontproperties=FONT, color=muted, fontsize=8)
    ax_src.set_xlabel("百万人", color=muted, fontsize=8, fontproperties=FONT)
    ax_src.invert_yaxis()
    for yv, val in zip(y_pos, src_top["value"][::-1] / 1e6):
        ax_src.text(val + 0.05, yv, f"{val:.1f}", va="center", color=text_c, fontsize=8, fontproperties=FONT)
    ax_src.set_xlim(0, max(src_top["value"] / 1e6) * 1.18)

    # 4.4 Top 庇护国
    ax_as = fig.add_axes(positions["asylum"])
    style_panel(ax_as, "04  主要庇护国（入边权重）")
    as_top = edges.groupby(["coa_iso", "target_cn"])["value"].sum().sort_values(ascending=False).head(10).reset_index()
    y_pos = np.arange(len(as_top))[::-1]
    colors_as = [cyan if i < 4 else "#4e8dad" for i in range(len(as_top))]
    ax_as.barh(y_pos, as_top["value"][::-1] / 1e6, color=colors_as[::-1], alpha=0.86)
    ax_as.set_yticks(y_pos)
    as_labels = [f"{a} {b}" for a, b in zip(as_top["coa_iso"], as_top["target_cn"])]
    ax_as.set_yticklabels(as_labels[::-1], fontproperties=FONT, color=muted, fontsize=8)
    ax_as.set_xlabel("百万人", color=muted, fontsize=8, fontproperties=FONT)
    ax_as.invert_yaxis()
    for yv, val in zip(y_pos, as_top["value"][::-1] / 1e6):
        ax_as.text(val + 0.04, yv, f"{val:.1f}", va="center", color=text_c, fontsize=8, fontproperties=FONT)
    ax_as.set_xlim(0, max(as_top["value"] / 1e6) * 1.18)

    # 4.5 区域流向矩阵
    ax_mat = fig.add_axes(positions["matrix"])
    style_panel(ax_mat, "05  区域流向矩阵（来源区域 → 庇护区域）")
    region_order = ["非洲", "亚洲", "欧洲", "美洲", "大洋洲", "其他"]
    mat = edges.pivot_table(index="source_region", columns="target_region", values="value", aggfunc="sum", fill_value=0)
    mat = mat.reindex(index=region_order, columns=region_order, fill_value=0)
    arr = np.log10(mat.values + 1)
    im = ax_mat.imshow(arr, cmap=cmap_heat, aspect="auto")
    ax_mat.set_xticks(range(len(region_order))); ax_mat.set_yticks(range(len(region_order)))
    ax_mat.set_xticklabels(region_order, color=muted, fontsize=8, fontproperties=FONT)
    ax_mat.set_yticklabels(region_order, color=muted, fontsize=8, fontproperties=FONT)
    ax_mat.set_xlabel("庇护区域", color=muted, fontsize=8, fontproperties=FONT)
    ax_mat.set_ylabel("来源区域", color=muted, fontsize=8, fontproperties=FONT)
    for i in range(arr.shape[0]):
        for j in range(arr.shape[1]):
            if mat.values[i, j] > 0:
                ax_mat.text(j, i, fmt_people(mat.values[i, j]), ha="center", va="center",
                            color="#f2f8ff" if arr[i, j] > arr.max() * 0.50 else "#adc3cf",
                            fontsize=7, fontproperties=FONT)
    cax = fig.add_axes([positions["matrix"][0] + 0.238, positions["matrix"][1] + 0.030, 0.085, 0.012])
    cb = plt.colorbar(im, cax=cax, orientation="horizontal")
    cb.ax.tick_params(labelsize=6, colors=muted, length=0)
    cb.set_ticks([])
    cb.outline.set_edgecolor("#31546a")
    fig.text(positions["matrix"][0] + 0.240, positions["matrix"][1] + 0.047, "低    log权重    高", color=muted, fontsize=7, fontproperties=FONT)

    # 4.6 趋势
    ax_tr = fig.add_axes(positions["trend"])
    style_panel(ax_tr, "06  网络规模变化（2010–2022）")
    trend = pop.groupby("year").agg(total_value=("value", "sum"), edge_count=("value", "size")).reset_index()
    ax_tr.plot(trend["year"], trend["total_value"] / 1e6, marker="o", markersize=3.5, color=gold, lw=2.0, label="总权重")
    ax_tr.set_ylabel("总权重（百万人）", color=gold, fontsize=8, fontproperties=FONT)
    ax_tr.tick_params(axis="y", labelcolor=gold)
    ax_tr2 = ax_tr.twinx()
    ax_tr2.plot(trend["year"], trend["edge_count"], marker="o", markersize=3.0, color=cyan, lw=1.7, label="边数量")
    ax_tr2.set_ylabel("")
    ax_tr2.tick_params(axis="y", labelright=False, right=False)
    for sp in ax_tr2.spines.values(): sp.set_color("#24475c")
    ax_tr.set_xlabel("年份", color=muted, fontsize=8, fontproperties=FONT)
    ax_tr.set_xticks([2010, 2013, 2016, 2019, 2022])
    ax_tr.text(0.04, 0.88, "黄线=总权重；蓝线=有效边数。排除本国境内与未知/无国籍项",
               transform=ax_tr.transAxes, color="#7fa2b5", fontsize=7.2, fontproperties=FONT)
    ax_tr.text(0.78, 0.17, f"2022边数：{int(trend.loc[trend['year']==YEAR, 'edge_count'].iloc[0]):,}",
               transform=ax_tr.transAxes, color=cyan, fontsize=7.2, fontproperties=FONT_BOLD)

    # 4.7 社区结构网络缩略图
    ax_com = fig.add_axes(positions["community"])
    style_panel(ax_com, "07  核心网络社区结构（抽象布局）")
    com_edges = edges[edges["value"] >= COMMUNITY_EDGE_MIN]
    Gu = nx.Graph()
    for r in com_edges.itertuples(index=False):
        Gu.add_edge(r.coo_iso, r.coa_iso, weight=float(r.value))
    # 保留最大连通与高强度节点，避免图过密
    strength = dict(Gu.degree(weight="weight"))
    keep_nodes = set(pd.Series(strength).sort_values(ascending=False).head(70).index)
    Gu = Gu.subgraph(keep_nodes).copy()
    communities = list(nx.algorithms.community.greedy_modularity_communities(Gu, weight="weight")) if Gu.number_of_edges() else []
    com_map = {}
    for i, com in enumerate(communities):
        for node in com:
            com_map[node] = i
    pos = nx.spring_layout(Gu, weight="weight", seed=SEED, k=0.45, iterations=220)
    # 绘边
    edge_widths = [0.2 + 2.2 * np.log10(Gu[u][v]["weight"] + 1) / np.log10(max(1, max(nx.get_edge_attributes(Gu, "weight").values())) + 1) for u, v in Gu.edges()]
    nx.draw_networkx_edges(Gu, pos, ax=ax_com, edge_color="#708fa0", width=edge_widths, alpha=0.20, arrows=False)
    # 绘点
    node_sizes = [35 + 260 * np.log10(strength.get(n, 1) + 1) / np.log10(max(strength.values()) + 1) for n in Gu.nodes()]
    node_cols = [community_colors[com_map.get(n, 0) % len(community_colors)] for n in Gu.nodes()]
    nx.draw_networkx_nodes(Gu, pos, ax=ax_com, node_size=node_sizes, node_color=node_cols,
                           alpha=0.88, linewidths=0.5, edgecolors="#ffffff")
    top_label_nodes = pd.Series(strength).sort_values(ascending=False).head(12).index
    for n in top_label_nodes:
        if n in pos:
            ax_com.text(pos[n][0] + 0.02, pos[n][1] + 0.02, CN_NAMES.get(n, n), fontsize=7, color=text_c,
                        fontproperties=FONT_BOLD, path_effects=[pe.withStroke(linewidth=1.8, foreground=panel_bg)])
    ax_com.text(0.02, 0.04, f"社区数：{len(communities)}    边阈值：≥{COMMUNITY_EDGE_MIN:,}",
                transform=ax_com.transAxes, color=muted, fontsize=7.5, fontproperties=FONT)
    ax_com.set_xticks([]); ax_com.set_yticks([])

    # 4.8 Top corridors + 分析解读
    ax_cor = fig.add_axes(positions["corridor"])
    ax_cor.set_facecolor(panel_bg)
    for sp in ax_cor.spines.values():
        sp.set_color("#24475c"); sp.set_linewidth(0.9)
    ax_cor.set_xticks([]); ax_cor.set_yticks([])
    ax_cor.set_title("08  关键走廊与结构解读", loc="left", pad=8, color=text_c, fontsize=11, fontproperties=FONT_BOLD)
    top_corr = edges.nlargest(7, "value").copy()
    y0 = 0.84
    for i, r in enumerate(top_corr.itertuples(index=False), start=1):
        y = y0 - (i - 1) * 0.085
        ax_cor.text(0.03, y, f"{i}", transform=ax_cor.transAxes, color=gold, fontsize=8.5, fontproperties=FONT_BOLD,
                    bbox=dict(boxstyle="circle,pad=0.18", facecolor="#102b3c", edgecolor="#3b6578", linewidth=0.5))
        ax_cor.text(0.095, y, f"{r.source_cn} → {r.target_cn}", transform=ax_cor.transAxes,
                    color=text_c, fontsize=8.5, fontproperties=FONT_BOLD)
        ax_cor.text(0.74, y, fmt_people(r.value), transform=ax_cor.transAxes,
                    color="#f8d47a", fontsize=8.5, fontproperties=FONT_BOLD, ha="right")
        ax_cor.plot([0.095, 0.70], [y - 0.025, y - 0.025], transform=ax_cor.transAxes,
                    color=cmap_flow((np.log10(r.value) - np.log10(top_corr["value"].min())) / (np.log10(top_corr["value"].max()) - np.log10(top_corr["value"].min()) + 1e-9)),
                    lw=1.6 + 4.0 * r.value / top_corr["value"].max(), alpha=0.80)
    explanation = (
        "结构特征：网络并非均匀扩散，而是由少数高权重走廊主导。"
        "亚洲—欧洲、拉美区域、非洲邻国之间形成三类显著模块；"
        "右上角色图中的双重节点体现来源与接收身份的叠加。"
    )
    ax_cor.text(0.03, 0.06, textwrap.fill(explanation, 34), transform=ax_cor.transAxes,
                color="#a9c6d7", fontsize=8.2, linespacing=1.6, fontproperties=FONT)

    # 页脚
    fig.text(0.030, 0.030, "说明：本图由 Python（Pandas + NetworkX + Matplotlib/Basemap）读取公开 CSV 后计算绘制；权重为年末存量口径，不代表实时移动路径。",
             color="#6f8fa0", fontsize=8.5, fontproperties=FONT)
    fig.text(0.970, 0.030, "网络可视化课程作业  |  数据主题：跨境保护关系网络", ha="right",
             color="#6f8fa0", fontsize=8.5, fontproperties=FONT)

    fig.savefig(OUT_PNG, dpi=180, facecolor=bg, bbox_inches="tight", pad_inches=0.08)
    plt.close(fig)

# ========== 5. 文档 ==========
def make_docx() -> None:
    text = "本作品基于UNHCR公开数据构建2022年跨境保护网络。节点为国家，边为来源国至庇护国关系，结合地图流线、角色矩阵、区域热力与社区图分析主要来源、接收枢纽及区域结构。"
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(54)
    sec.left_margin = Pt(64)
    sec.right_margin = Pt(64)
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    r.font.size = Pt(12)
    doc.save(OUT_DOCX)


def main() -> None:
    ensure_data()
    pop, countries, edges, G = read_data()
    make_visualization(pop, countries, edges, G)
    make_docx()
    print(f"Done: {OUT_PNG}")
    print(f"Done: {OUT_DOCX}")


if __name__ == "__main__":
    main()
